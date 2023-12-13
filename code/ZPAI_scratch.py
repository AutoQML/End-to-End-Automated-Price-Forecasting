import matplotlib
import math
import numpy as np
from matplotlib import pyplot as plt

from pathlib import Path

from docx.shared import Inches

from ZPAI_document_results_utility import get_min_values, get_max_values

# nn = [
#     [0.23], # basic
#     [0.18], # extension
#     [0.18], # XE
#     [0.19], # location
#     [0.19], # extension-XE
#     [0.19], # extension-location
#     [0.2], # XE-location
#     [0.19] # extension-XE-location
# ]

def modify_legend(ax):
    handles, labels_txt = ax.get_legend_handles_labels()

    # Capitalize the first letters
    labels_txt = [label.capitalize() for label in labels_txt]
    # replace 'manual' string by the actual algorithm 'Rand. Forest'
    labels_txt = [label.replace('Manual', 'Random Forest') for label in labels_txt]
    # replace 'nn' string by the actual algorithm 'Neural Networkt'
    labels_txt = [label.replace('Nn', 'MLP') for label in labels_txt]
    # replace 'Autogluon' string by 'AutoGluon' string
    labels_txt = [label.replace('Autogluon', 'AutoGluon') for label in labels_txt]
    # replace 'Autosklearn' string by 'auto-sklearn' string
    labels_txt = [label.replace('Autosklearn', 'auto-sklearn') for label in labels_txt]
    # replace 'Flaml' string by 'FLAML' string
    labels_txt = [label.replace('Flaml', 'FLAML') for label in labels_txt]
    # replace 'Autokeras' string by 'AutoKeras' string
    labels_txt = [label.replace('Autokeras', 'AutoKeras') for label in labels_txt]

    return handles, labels_txt

def plot_dataset_performance(values: np.ndarray, labels: list, datasets: list, CHART_PATH: Path, CHART_PDF_PATH: Path, min_value: float, max_value: float, dataset: str, document, picture_size):

    MIN_VAL = min_value
    MAX_VAL = max_value
    dataset = dataset

    # values: frameworks x dataset x repetitions
    rows = values.shape[1]
    # a4_size = (8.27, 1.0 * rows)
    a4_size = (13.9, 1.6 * rows)
    # a4_size = (13.5, 2.2 * rows)

    fig, ax = plt.subplots(1, 1, gridspec_kw={'wspace': 0.01, 'hspace': 0})
    fig.set_size_inches(a4_size)

    plt.subplots_adjust(left=0.20)

    ax.set_frame_on(False)
    ax.grid(True, linewidth=0.5, alpha=0.25, color='black')
    ax.set_axisbelow(True)
    ax.set_yticks(np.arange(rows))
    # ax.set_xticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0])
    # set dynamic xticks steps
    xticks_step_size_range = (MAX_VAL - MIN_VAL)
    xticks_step_size = 0.05 if xticks_step_size_range > 0.2 else 0.01
    # xticks_step_size = 0.02 if xticks_step_size_range > 0.2 else 0.01  # for smaller grids 
    ax.set_xticks(np.arange(0.0, 1.1, xticks_step_size))

    ax.tick_params(axis=u'both', which=u'both', length=0)
    ax.set_yticklabels(datasets)
    # ax.set_xlim([-0.05, 1.05])
    # ax.set_xlim([ 0.145, 0.255])
    ax.set_xlim([MIN_VAL - xticks_step_size, MAX_VAL + xticks_step_size])
    # ax.set_xlim([0.121, 0.33]) # used for adaptation of the 2910 dataset
    ax.tick_params(axis='both', which='major', labelsize=14)

    # y_offsets = np.linspace(-0.15, 0.15, values.shape[0])
    y_offsets = np.linspace(-0.00, 0.00, values.shape[0])
    for idx in range(values.shape[0]):
        mean = values[idx].mean(axis=1)
        mean_y = np.arange(rows) + y_offsets[idx]
        x = values[idx]
        y = np.arange(x.shape[0]).repeat(x.shape[1]).reshape(x.shape) + y_offsets[idx]

        # display singe measurements and mean value of all measurements
        color = next(ax._get_lines.prop_cycler)
        ax.scatter(x, y, s=(matplotlib.rcParams['lines.markersize'] ** 2) * 0.33, alpha=0.45, linewidths=0, **color) # display dots for sigle measurements
        ax.scatter(mean, mean_y, marker='d', s=(matplotlib.rcParams['lines.markersize'] ** 2.7) * .9, label=labels[idx], **color) # display diamonds 

    graphics_title = f"MAPE-scores for {dataset}"
    # plt.title(graphics_title) # uncomment for publication

    ################
    # modify y axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_yticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Add 'basic-subset' to all labels except tha basic-subset label
    labels = [label.replace(label, 'basic-subset +\n{}'.format(label)) if label !='basic-subset' else 'basic-subset' for label in labels]

    labels = ['']

    ax.set_yticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = modify_legend(ax)

    fig.subplots_adjust(bottom=0.3 / rows)
    fig.set_figheight(4)
    fig.legend(handles, labels_txt, ncol=len(labels_txt), loc='lower center', borderaxespad=1.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=6, loc='lower center', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc='right', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc=(0.85, 0.67), borderaxespad=0.5, fontsize=13)

    plt.xlabel("Mean absolute percentage error [MAPE]", fontsize=16)
    plt.ylabel(" ", fontsize=16)
    # save grafics
    plt.savefig(CHART_PATH)
    plt.savefig(CHART_PDF_PATH)
    # integrate grafics into dokumentation
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()

def plot_feature_performance(score_result_df, score, CHART_PATH, CHART_PDF_PATH, max_value, min_value, dataset, document, picture_size):
    max_value =  max_value
    min_value = min_value
    dataset = dataset

    feature_set_names = list(score_result_df.columns.values) # get names of the feature sets
    # print(feature_set_names)
    algorithm_names = list(score_result_df.index.values) # get names of the algorithms
    # print(algorithm_names)

    # add a data row for each item
    # for i in range(len(algorithm_names)):
    result_values = []
    for i, val in enumerate(algorithm_names):
        # print(val)
        # list(val)
        val = []
        # for j in range(len(r2_extended_column_names)):
        for j, value in enumerate(feature_set_names):
            # print(score_result_df.iat[i, j])
            val.append(score_result_df.iat[i, j])

        # print(val)
        result_values.append(val)
        
    # print(result_values)

    plot_dataset_performance(
        np.asarray(result_values),
        algorithm_names,
        feature_set_names,
        CHART_PATH,
        CHART_PDF_PATH,
        min_value,
        max_value,
        dataset,
        document, 
        picture_size
)

###########################
# Training duration plot
###########################
    
def subfunc_plot_training_duration(values: np.ndarray, labels: list, datasets: list, CHART_PATH: Path, CHART_PDF_PATH: Path, min_value: float, max_value: float, dataset: str, document, picture_size):

    MIN_VAL = min_value
    MAX_VAL = max_value
    dataset = dataset

    rounded_max_val = math.ceil(MAX_VAL / 1000) * 1000
    rounded_min_val = math.floor(MIN_VAL / 1000) * 1000


    # values: frameworks x dataset x repetitions
    rows = values.shape[1]
    # a4_size = (8.27, 1.0 * rows)
    a4_size = (13.9, 1.6 * rows)

    fig, ax = plt.subplots(1, 1, gridspec_kw={'wspace': 0.01, 'hspace': 0})
    fig.set_size_inches(a4_size)

    plt.subplots_adjust(left=0.20)

    ax.set_frame_on(False)
    ax.grid(True, linewidth=0.5, alpha=0.25, color='black')
    ax.set_axisbelow(True)
    ax.set_yticks(np.arange(rows))
    # ax.set_xticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0])
    ax.set_xticks(np.arange(rounded_min_val, rounded_max_val, 500))
    ax.tick_params(axis=u'both', which=u'both', length=0)
    ax.set_yticklabels(datasets)
    # ax.set_xlim([-0.05, 1.05])
    # ax.set_xlim([ -100.0, 1700])
    ax.set_xlim([MIN_VAL - 100.0, MAX_VAL + 200.0])
    ax.tick_params(axis='both', which='major', labelsize=14)

    y_offsets = np.linspace(-0.15, 0.15, values.shape[0])
    for idx in range(values.shape[0]):
        mean = values[idx].mean(axis=1)
        mean_y = np.arange(rows) + y_offsets[idx]
        x = values[idx]
        y = np.arange(x.shape[0]).repeat(x.shape[1]).reshape(x.shape) + y_offsets[idx]

        # display singe measurements and mean value of all measurements
        color = next(ax._get_lines.prop_cycler)
        ax.scatter(x, y, s=(matplotlib.rcParams['lines.markersize'] ** 2) * 0.33, alpha=0.45, linewidths=0, **color)
        ax.scatter(mean, mean_y, marker='d', s=(matplotlib.rcParams['lines.markersize'] ** 2.7) * .9, label=labels[idx],
                   **color)

    graphics_title = f"MAPE-scores for {dataset}"
    # plt.title(graphics_title) # uncomment for publication

    ################
    # modify y axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_yticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Add 'basic-subset' to all labels except tha basic-subset label
    labels = [label.replace(label, 'basic-subset +\n{}'.format(label)) if label !='basic-subset' else 'basic-subset' for label in labels]

    ax.set_yticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = modify_legend(ax)

    fig.subplots_adjust(bottom=0.8 / rows)
    fig.legend(handles, labels_txt, ncol=len(labels_txt), loc='lower center', borderaxespad=1.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=3, loc='lower center', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc='right', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc=(0.85, 0.67), borderaxespad=0.5, fontsize=13)

    plt.xlabel("Training time (sec)", fontsize=16)
    plt.ylabel("Feature combination", fontsize=16)
    # save grafics
    plt.savefig(CHART_PATH)
    plt.savefig(CHART_PDF_PATH)
    # integrate grafics into dokumentation
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()
    
def plot_training_duration(score_result_df, score, CHART_PATH, CHART_PDF_PATH, max_value, min_value, dataset, document, picture_size):
    max_value =  max_value
    min_value = min_value
    dataset = dataset

    feature_set_names = list(score_result_df.columns.values) # get names of the feature sets
    # print(feature_set_names)
    algorithm_names = list(score_result_df.index.values) # get names of the algorithms
    # print(algorithm_names)

    # add a data row for each item
    # for i in range(len(algorithm_names)):
    result_values = []
    for i, val in enumerate(algorithm_names):
        # print(val)
        # list(val)
        val = []
        # for j in range(len(r2_extended_column_names)):
        for j, value in enumerate(feature_set_names):
            # print(score_result_df.iat[i, j])
            val.append(score_result_df.iat[i, j])

        # print(val)
        result_values.append(val)
        
    # print(result_values)

    subfunc_plot_training_duration(
        np.asarray(result_values),
        algorithm_names,
        feature_set_names,
        CHART_PATH,
        CHART_PDF_PATH,
        min_value,
        max_value,
        dataset,
        document, 
        picture_size
)

###########################
# Testing duration plot
###########################
    
def subfunc_plot_testing_duration(values: np.ndarray, labels: list, datasets: list, CHART_PATH: Path, CHART_PDF_PATH: Path, min_value: float, max_value: float, dataset: str, document, picture_size):

    MIN_VAL = min_value
    MAX_VAL = max_value
    dataset = dataset

    # values: frameworks x dataset x repetitions
    rows = values.shape[1]
    # a4_size = (8.27, 1.0 * rows)
    a4_size = (13.5, 1.6 * rows)

    fig, ax = plt.subplots(1, 1, gridspec_kw={'wspace': 0.01, 'hspace': 0})
    fig.set_size_inches(a4_size)

    plt.subplots_adjust(left=0.20)

    ax.set_frame_on(False)
    ax.grid(True, linewidth=0.5, alpha=0.25, color='black')
    ax.set_axisbelow(True)
    ax.set_yticks(np.arange(rows))
    ax.set_xticks(np.arange(0.0, MAX_VAL, 0.002))
    ax.tick_params(axis=u'both', which=u'both', length=0)
    ax.set_yticklabels(datasets)
    # ax.set_xlim([-0.05, 1.05])
    # ax.set_xlim([ -100.0, 1700])
    ax.set_xlim([MIN_VAL - 0.00015, MAX_VAL + 0.00015])
    ax.tick_params(axis='both', which='major', labelsize=14)

    y_offsets = np.linspace(-0.15, 0.15, values.shape[0])
    for idx in range(values.shape[0]):
        mean = values[idx].mean(axis=1)
        mean_y = np.arange(rows) + y_offsets[idx]
        x = values[idx]
        y = np.arange(x.shape[0]).repeat(x.shape[1]).reshape(x.shape) + y_offsets[idx]

        # display singe measurements and mean value of all measurements
        color = next(ax._get_lines.prop_cycler)
        ax.scatter(x, y, s=(matplotlib.rcParams['lines.markersize'] ** 2) * 0.33, alpha=0.45, linewidths=0, **color)
        ax.scatter(mean, mean_y, marker='d', s=(matplotlib.rcParams['lines.markersize'] ** 2.7) * .9, label=labels[idx],
                   **color)

    graphics_title = f"MAPE-scores for {dataset}"
    # plt.title(graphics_title) # uncomment for publication

    ################
    # modify y axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_yticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Add 'basic-subset' to all labels except tha basic-subset label
    labels = [label.replace(label, 'basic-subset +\n{}'.format(label)) if label !='basic-subset' else 'basic-subset' for label in labels]

    ax.set_yticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = modify_legend(ax)

    fig.subplots_adjust(bottom=0.8 / rows)
    fig.legend(handles, labels_txt, ncol=len(labels_txt), loc='lower center', borderaxespad=1.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=3, loc='lower center', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc='right', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc=(0.85, 0.67), borderaxespad=0.5, fontsize=13)

    plt.xlabel("Prediction time (sec)", fontsize=16)
    plt.ylabel("Feature combination", fontsize=16)
    # save grafics
    plt.savefig(CHART_PATH)
    plt.savefig(CHART_PDF_PATH)
    # integrate grafics into dokumentation
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()
    
def plot_testing_duration(score_result_df, score, CHART_PATH, CHART_PDF_PATH, max_value, min_value, dataset, document, picture_size):
    max_value =  max_value
    min_value = min_value
    dataset = dataset

    feature_set_names = list(score_result_df.columns.values) # get names of the feature sets
    # print(feature_set_names)
    algorithm_names = list(score_result_df.index.values) # get names of the algorithms
    # print(algorithm_names)

    # add a data row for each item
    # for i in range(len(algorithm_names)):
    result_values = []
    for i, val in enumerate(algorithm_names):
        # print(val)
        # list(val)
        val = []
        # for j in range(len(r2_extended_column_names)):
        for j, value in enumerate(feature_set_names):
            # print(score_result_df.iat[i, j])
            val.append(score_result_df.iat[i, j])

        # print(val)
        result_values.append(val)
        
    # print(result_values)

    subfunc_plot_testing_duration(
        np.asarray(result_values),
        algorithm_names,
        feature_set_names,
        CHART_PATH,
        CHART_PDF_PATH,
        min_value,
        max_value,
        dataset,
        document, 
        picture_size
)


###########################
# MES plot
###########################
    
def subfunc_plot_mev(values: np.ndarray, labels: list, datasets: list, CHART_PATH: Path, CHART_PDF_PATH: Path, min_value: float, max_value: float, dataset: str, document, picture_size):

    MIN_VAL = min_value
    MAX_VAL = max_value
    dataset = dataset

    # values: frameworks x dataset x repetitions
    rows = values.shape[1]
    # a4_size = (8.27, 1.0 * rows)
    a4_size = (13.5, 1.6 * rows)

    fig, ax = plt.subplots(1, 1, gridspec_kw={'wspace': 0.01, 'hspace': 0})
    fig.set_size_inches(a4_size)

    plt.subplots_adjust(left=0.20)

    ax.set_frame_on(False)
    ax.grid(True, linewidth=0.5, alpha=0.25, color='black')
    ax.set_axisbelow(True)
    ax.set_yticks(np.arange(rows))
    # ax.set_xticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0, 1.2])
    ax.set_xticks(np.arange(0.0, 1.2, 0.1))
    ax.tick_params(axis=u'both', which=u'both', length=0)
    ax.set_yticklabels(datasets)
    ax.set_xlim([MIN_VAL - 0.05, MAX_VAL + 0.05])
    ax.tick_params(axis='both', which='major', labelsize=14)

    y_offsets = np.linspace(-0.15, 0.15, values.shape[0])
    for idx in range(values.shape[0]):
        mean = values[idx].mean(axis=1)
        mean_y = np.arange(rows) + y_offsets[idx]
        # x = values[idx]
        # y = np.arange(x.shape[0]).repeat(x.shape[1]).reshape(x.shape) + y_offsets[idx]

        # display singe measurements and mean value of all measurements
        color = next(ax._get_lines.prop_cycler)
        # ax.scatter(x, y, s=(matplotlib.rcParams['lines.markersize'] ** 2) * 0.33, alpha=0.45, linewidths=0, **color)
        ax.scatter(mean, mean_y, marker='d', s=(matplotlib.rcParams['lines.markersize'] ** 2.7) * .9, label=labels[idx],
                   **color)

    graphics_title = f"MEV-scores for {dataset}"
    # plt.title(graphics_title) # uncomment for publication

    ################
    # modify y axis labels
    #################
    # Extract the label names 
    labels = [item.get_text() for item in ax.get_yticklabels()]
    # Capitalize the first letters of the label names
    # labels = [label.capitalize() for label in labels]
    # Add 'basic-subset' to all labels except tha basic-subset label
    labels = [label.replace(label, 'basic-subset +\n{}'.format(label)) if label !='basic-subset' else 'basic-subset' for label in labels]

    ax.set_yticklabels(labels)

    ################
    # modify legend
    #################
    handles, labels_txt = modify_legend(ax)

    fig.subplots_adjust(bottom=0.8 / rows)
    fig.legend(handles, labels_txt, ncol=len(labels_txt), loc='lower center', borderaxespad=1.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=3, loc='lower center', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc='right', borderaxespad=0.5, fontsize=13)
    # fig.legend(handles, labels_txt, ncol=1, loc=(0.85, 0.67), borderaxespad=0.5, fontsize=13)

    plt.xlabel("Method evaluation score [MES]", fontsize=16)
    plt.ylabel("Feature combination", fontsize=16)
    # save grafics
    plt.savefig(CHART_PATH)
    plt.savefig(CHART_PDF_PATH)
    # integrate grafics into dokumentation
    document.add_picture(str(CHART_PATH), width=Inches(picture_size))

    plt.close()
    
def plot_mev(score_result_df, CHART_PATH, CHART_PDF_PATH, max_value, min_value, dataset, document, picture_size):
    max_value =  max_value
    min_value = min_value
    dataset = dataset

    feature_set_names = list(score_result_df.columns.values) # get names of the feature sets
    # print(feature_set_names)
    algorithm_names = list(score_result_df.index.values) # get names of the algorithms
    # print(algorithm_names)

    # add a data row for each item
    # for i in range(len(algorithm_names)):
    result_values = []
    for i, val in enumerate(algorithm_names):
        # print('val: ', val)
        # list(val)
        val = []
        # for j in range(len(r2_extended_column_names)):
        for j, value in enumerate(feature_set_names):
            # print(score_result_df.iat[i, j])
            val.append([score_result_df.iat[i, j]])

        # print(val)
        result_values.append(val)
        
    # print(result_values)

    subfunc_plot_mev(
        np.asarray(result_values),
        algorithm_names,
        feature_set_names,
        CHART_PATH,
        CHART_PDF_PATH,
        min_value,
        max_value,
        dataset,
        document, 
        picture_size
)